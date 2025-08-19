"""
File content extraction utilities for attachments (PDF, DOCX, TXT).
 
Dependencies (install if needed):
- pypdf         -> pip install pypdf
- python-docx   -> pip install python-docx
"""
 
from io import BytesIO
from typing import Optional, List, Dict
from loguru import logger
import pypdf  # modern package
import PyPDF2
try:
    from docx import Document  # comes from python-docx package
except ImportError:
    Document = None
 
 
def extract_text_from_pdf(file_bytes: bytes) -> str:
    """
    Extract text from a PDF file using pypdf.
   
    Args:
        file_bytes: Raw PDF file bytes
       
    Returns:
        Extracted text (empty string on failure)
       
    Note: Scanned/image-only PDFs will not yield text (no OCR here).
    """
    try:
        try:
            reader = pypdf.PdfReader(BytesIO(file_bytes))
            pages = reader.pages
            get_text = lambda p: p.extract_text() or ""
        except ImportError:
            # Fallback: older PyPDF2 if available
            try:
                reader = PyPDF2.PdfReader(BytesIO(file_bytes))
                pages = reader.pages
                get_text = lambda p: p.extract_text() or ""
            except ImportError:
                logger.error("No PDF library found. Install: pip install pypdf")
                return ""
 
        texts = []
        for page_num, page in enumerate(pages):
            try:
                page_text = get_text(page)
                if page_text:
                    texts.append(f"Page {page_num + 1}: {page_text}")
            except Exception as exc:
                logger.debug(f"PDF page {page_num + 1} extract error: {exc}")
       
        extracted_text = "\n\n".join(texts)
        logger.info(f"PDF extraction: {len(pages)} pages, {len(extracted_text)} characters")
        return extracted_text
       
    except Exception as exc:
        logger.error(f"PDF extraction failed: {exc}")
        return ""
 
 
def extract_text_from_docx(file_bytes: bytes) -> str:
    """
    Extract text from a DOCX file using python-docx.
   
    Args:
        file_bytes: Raw DOCX file bytes
       
    Returns:
        Extracted text (empty string on failure)
    """
    if Document is None:
        logger.error("python-docx not found. Install: pip install python-docx")
        return ""
        
    try:
        doc = Document(BytesIO(file_bytes))
       
        # Extract text from paragraphs
        paragraphs = []
        for para in doc.paragraphs:
            if para.text.strip():
                paragraphs.append(para.text.strip())
       
        # Extract text from tables
        tables_text = []
        for table in doc.tables:
            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    if cell.text.strip():
                        row_text.append(cell.text.strip())
                if row_text:
                    tables_text.append(" | ".join(row_text))
       
        # Combine all text
        all_text = []
        if paragraphs:
            all_text.append("Paragraphs:\n" + "\n".join(paragraphs))
        if tables_text:
            all_text.append("Tables:\n" + "\n".join(tables_text))
       
        extracted_text = "\n\n".join(all_text)
        logger.info(f"DOCX extraction: {len(paragraphs)} paragraphs, {len(doc.tables)} tables, {len(extracted_text)} characters")
        return extracted_text
       
    except ImportError:
        logger.error("python-docx not found. Install: pip install python-docx")
        return ""
    except Exception as exc:
        logger.error(f"DOCX extraction failed: {exc}")
        return ""
 
 
def extract_text_from_txt(file_bytes: bytes, encoding: str = "utf-8") -> str:
    """
    Extract text from a plain text file.
   
    Args:
        file_bytes: Raw text file bytes
        encoding: Text encoding to use (default: utf-8)
       
    Returns:
        Extracted text (empty string on failure)
    """
    try:
        # Try the specified encoding first
        try:
            text = file_bytes.decode(encoding)
        except UnicodeDecodeError:
            # Fallback to common encodings
            for fallback_encoding in ["utf-8", "latin-1", "cp1252", "iso-8859-1"]:
                try:
                    text = file_bytes.decode(fallback_encoding, errors="ignore")
                    logger.info(f"Used fallback encoding: {fallback_encoding}")
                    break
                except UnicodeDecodeError:
                    continue
            else:
                logger.error("Failed to decode text file with any encoding")
                return ""
       
        # Clean up the text
        lines = text.split('\n')
        cleaned_lines = []
        for line in lines:
            line = line.strip()
            if line:  # Remove empty lines
                cleaned_lines.append(line)
       
        extracted_text = "\n".join(cleaned_lines)
        logger.info(f"TXT extraction: {len(cleaned_lines)} lines, {len(extracted_text)} characters")
        return extracted_text
       
    except Exception as exc:
        logger.error(f"TXT extraction failed: {exc}")
        return ""


def extract_text_from_attachments(attachments_data: List[Dict]) -> str:
    """
    Extract text from multiple attachments and combine them.
    
    Args:
        attachments_data: List of attachment dictionaries with 'Body' and 'Name' fields
        
    Returns:
        Combined text from all attachments
    """
    if not attachments_data:
        return ""
    
    all_texts = []
    
    for attachment in attachments_data:
        try:
            file_name = attachment.get('Name', 'unknown')
            file_body = attachment.get('Body', '')
            
            if not file_body:
                logger.warning(f"Empty body for attachment: {file_name}")
                continue
            
            # Get file extension from name or content type
            if '.' in file_name:
                file_extension = file_name.split('.')[-1].lower()
            elif attachment.get('ContentType'):
                # Map ContentType to file extension
                content_type = attachment.get('ContentType', '').lower()
                if 'pdf' in content_type:
                    file_extension = 'pdf'
                elif 'word' in content_type or 'docx' in content_type:
                    file_extension = 'docx'
                elif 'text' in content_type or 'plain' in content_type:
                    file_extension = 'txt'
                else:
                    logger.warning(f"Unsupported content type: {content_type} for file: {file_name}")
                    continue
            else:
                logger.warning(f"No file extension or content type found for: {file_name}")
                continue
            
            # Extract text based on file type
            if file_extension == 'pdf':
                extracted_text = extract_text_from_pdf(file_body)
            elif file_extension in ['docx', 'doc']:
                extracted_text = extract_text_from_docx(file_body)
            elif file_extension == 'txt':
                extracted_text = extract_text_from_txt(file_body)
            else:
                logger.warning(f"Unsupported file extension: {file_extension}")
                continue
            
            if extracted_text:
                all_texts.append(f"Attachment: {file_name}\n{extracted_text}")
                logger.info(f"Successfully extracted text from {file_name}: {len(extracted_text)} characters")
            else:
                logger.warning(f"No text extracted from {file_name}")
                
        except Exception as exc:
            logger.error(f"Error processing attachment {attachment.get('Name', 'unknown')}: {exc}")
            continue
    
    combined_text = "\n\n---\n\n".join(all_texts)
    logger.info(f"Total attachment text extracted: {len(combined_text)} characters from {len(attachments_data)} attachments")
    
    return combined_text


def test_extractors():
    """Test all extractors with sample files."""
    print("Testing file extractors...")
   
    # Test with empty bytes
    assert extract_text_from_pdf(b"") == ""
    assert extract_text_from_docx(b"") == ""
    assert extract_text_from_txt(b"") == ""
   
    print("All extractors handle empty input correctly")
    print("Note: Install pypdf and python-docx to test with real files")


if __name__ == "__main__":
    test_extractors()